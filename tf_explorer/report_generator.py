# -*- coding: utf-8 -*-
"""
TF-Explorer Research Report Generator
Programmatically constructs formatted, publication-grade Microsoft Word reports
containing active epigenetic search parameters, ChIP peaks, motifs, conservation,
primer designs, GTEx tissue median expression, and STRING interactomes.
"""

import os
from datetime import datetime
from io import BytesIO
import pandas as pd
from typing import Dict, List, Optional, Tuple, Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- XML Helper Functions for Premium Formatting ---

def set_cell_background(cell, fill_hex: str):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding (in dxa) for a table cell to ensure clean spacing."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_table_borders(table):
    """Adds a clean, double-lined border style or soft gray borders to tables."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    # Soft gray thin lines for inside grid, slightly thicker for top and bottom
    for border_name, color, sz, val in [
        ('top', '475569', '8', 'single'),
        ('bottom', '475569', '12', 'single'),
        ('insideH', 'cbd5e1', '4', 'single'),
        ('left', 'ffffff', '0', 'none'),
        ('right', 'ffffff', '0', 'none'),
        ('insideV', 'ffffff', '0', 'none')
    ]:
        node = OxmlElement(f'w:{border_name}')
        node.set(qn('w:val'), val)
        node.set(qn('w:sz'), sz)
        node.set(qn('w:space'), '0')
        node.set(qn('w:color'), color)
        tblBorders.append(node)
    tblPr.append(tblBorders)

def add_heading_styled(doc, text: str, level: int) -> Any:
    """Adds a heading with specialized typography and color matching the premium theme."""
    h = doc.add_heading('', level=level)
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.bold = True
    
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(30, 41, 59)  # Slate 800 (#1e293b)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(71, 85, 105)  # Slate 600 (#475569)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
        
    return h

def generate_report_docx(
    gene_name: str,
    tf_list: List[str],
    promoter_up: int,
    promoter_down: int,
    genome: str,
    df_encode: pd.DataFrame,
    df_motifs: pd.DataFrame,
    df_cons: Optional[pd.DataFrame],
    df_synergy: Optional[pd.DataFrame],
    df_gtex: Optional[pd.DataFrame],
    string_res: Optional[Dict],
    primer_results: Optional[List[Tuple[str, Any]]],
    output_path: Optional[str] = None
) -> BytesIO:
    """
    Constructs the Word document, writes structured research reports, and returns
    it as an in-memory BytesIO stream or saves it to output_path.
    """
    
    doc = Document()
    
    # --- PAGE SETUP ---
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # --- TYPOGRAPHY DEFAULT STYLE ---
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(51, 65, 85) # Slate 700 (#334155)
    
    # --- TITLE & METADATA SECTION ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("TF-EXPLORER BIOLOGICAL RESEARCH SUMMARY REPORT")
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42) # Slate 900 (#0f172a)
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_subtitle.paragraph_format.space_after = Pt(18)
    run_sub = p_subtitle.add_run("Systems-Level Epigenetic Profiling & Promoter Binding Analysis Suite")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
    
    # Metadata Block Table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_borders(meta_table)
    
    metadata_fields = [
        ("Target Gene:", gene_name),
        ("Associated Transcription Factors:", ", ".join(tf_list)),
        ("Promoter Coordinates Range:", f"TSS -{promoter_up} bp to TSS +{promoter_down} bp"),
        ("Active Genome Assembly:", f"{genome} (Homo sapiens)"),
        ("Analysis Timestamp:", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    ]
    
    for i, (field_lbl, field_val) in enumerate(metadata_fields):
        row = meta_table.rows[i]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.2)
        c1.width = Inches(4.3)
        
        # Style label
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(field_lbl)
        r0.bold = True
        r0.font.color.rgb = RGBColor(30, 41, 59)
        
        # Style value
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(2)
        p1.add_run(field_val)
        
        # Cell margins
        set_cell_margins(c0, top=60, bottom=60, left=100, right=100)
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # --- 1. EXECUTIVE SUMMARY ---
    add_heading_styled(doc, "1. Executive Summary", level=1)
    
    sum_text = (
        f"This report presents an integrated chromatin and regulatory transcription factor binding analysis for the target gene {gene_name}. "
        f"By cross-referencing functional ENCODE ChIP-Seq raw binding peak profiles, JASPAR transcription factor binding sequence motifs, "
        f"and evolutionary conservation constraint maps, TF-Explorer has isolated functional regulatory zones inside the promoter sequence. "
        f"Additionally, median tissue-specific expression indexes and direct macromolecular networks are documented below to map overall systemic coregulation."
    )
    doc.add_paragraph(sum_text)
    
    # --- 2. ENCODE ChIP-SEQ OVERLAPS SUMMARY ---
    add_heading_styled(doc, "2. ENCODE ChIP-Seq Peak Overlaps", level=1)
    
    if not df_encode.empty:
        window_peaks = df_encode[
            (df_encode['distance_to_tss'] >= -promoter_up) & 
            (df_encode['distance_to_tss'] <= promoter_down)
        ].copy()
        
        strict_peaks = window_peaks[window_peaks['overlap'] == True] if 'overlap' in window_peaks.columns else window_peaks
        total_p = len(window_peaks)
        strict_p = len(strict_peaks)
        
        p_stats = doc.add_paragraph()
        p_stats.add_run("Binding Summary: ").bold = True
        p_stats.add_run(
            f"Within the promoter region, a total of {total_p} raw peak binding records were retrieved. "
            f"Of these, {strict_p} binding sites lie strictly within the high-occupancy promoter coordinate block. "
            f"These interactions demonstrate direct epigenetic occupancy across multiple cell-type configurations."
        )
        
        # Render a table showing top peaks
        if not strict_peaks.empty:
            add_heading_styled(doc, "Top Binding Peaks Detected", level=2)
            top_peaks = strict_peaks.sort_values(by=strict_peaks.columns[0]).head(10)
            
            headers = ["Experiment", "Biosample (Cell Line)", "Distance to TSS", "Signal Intensity"]
            peak_table = doc.add_table(rows=len(top_peaks) + 1, cols=4)
            peak_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_table_borders(peak_table)
            
            # Format Headers
            hdr_cells = peak_table.rows[0].cells
            for col_idx, text in enumerate(headers):
                set_cell_background(hdr_cells[col_idx], "1e293b")
                set_cell_margins(hdr_cells[col_idx], top=100, bottom=100, left=100, right=100)
                p = hdr_cells[col_idx].paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(text)
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                
            # Fill rows
            for r_idx, (_, row) in enumerate(top_peaks.iterrows()):
                cells = peak_table.rows[r_idx + 1].cells
                sig_val = row.get('signal', row.get('score', 0.0))
                if pd.isna(sig_val): sig_val = 0.0
                
                vals = [
                    str(row.get('experiment', 'N/A')),
                    str(row.get('biosample', 'N/A')),
                    f"{int(row.get('distance_to_tss', 0)):+,} bp",
                    f"{sig_val:.2f}"
                ]
                
                for col_idx, val in enumerate(vals):
                    set_cell_margins(cells[col_idx], top=60, bottom=60, left=100, right=100)
                    p = cells[col_idx].paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    p.add_run(val)
    else:
        doc.add_paragraph("No overlapping ENCODE ChIP-Seq peaks were retrieved or loaded for this promoter query.")
        
    # --- 3. MOTIF SYNERGY CLUSTERS ---
    add_heading_styled(doc, "3. Transcription Factor Motif Synergy", level=1)
    
    if df_synergy is not None and not df_synergy.empty:
        doc.add_paragraph(
            "TF Motif Synergy clusters represent regions where multiple distinct regulatory factors have predicted JASPAR binding motif sequences localized in tight clusters (within 100 bp). "
            "These regions represent high-confidence hotspots for synergistic combinatorial binding complexes."
        )
        
        syn_table = doc.add_table(rows=len(df_synergy) + 1, cols=4)
        syn_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_table_borders(syn_table)
        
        headers = ["Synergy Coordinates", "Relative TSS Range", "Co-occurring Factors", "Motifs Count"]
        hdr_cells = syn_table.rows[0].cells
        for col_idx, text in enumerate(headers):
            set_cell_background(hdr_cells[col_idx], "475569")
            set_cell_margins(hdr_cells[col_idx], top=100, bottom=100, left=100, right=100)
            p = hdr_cells[col_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
        for r_idx, row in df_synergy.iterrows():
            cells = syn_table.rows[r_idx + 1].cells
            start_rel = row['start'] - promoter_up
            end_rel = row['end'] - promoter_up
            
            vals = [
                f"{row['start']} to {row['end']} bp",
                f"{start_rel:+,} to {end_rel:+,} bp",
                str(row['tfs']),
                str(row['motifs_count'])
            ]
            
            for col_idx, val in enumerate(vals):
                set_cell_margins(cells[col_idx], top=60, bottom=60, left=100, right=100)
                p = cells[col_idx].paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                p.add_run(val)
    else:
        doc.add_paragraph("No clustered transcription factor motif synergy hotspots (clustered within 100 bp of different TFs) were detected in this analysis promoter window.")
        
    # --- 4. EVOLUTIONARY CONSERVATION PROFILE ---
    add_heading_styled(doc, "4. Evolutionary Conservation constraints", level=1)
    
    if df_cons is not None and not df_cons.empty:
        phast_mean = df_cons['phastCons'].mean()
        phast_max = df_cons['phastCons'].max()
        phylo_mean = df_cons['phyloP'].mean()
        
        doc.add_paragraph(
            f"Analysis of UCSC evolutionary constraints reveals base-by-base selective pressures across the promoter. "
            f"The promoter window has an average phastCons (100-way vertebrate alignment) score of {phast_mean:.4f} and a peak conservation score of {phast_max:.4f}. "
            f"An average phyloP substitution rate constraint score of {phylo_mean:.4f} was observed."
        )
        
        # High confidence conserved peaks table (phastCons > 0.8)
        if not df_encode.empty and 'max_conservation' in df_encode.columns:
            conserved_peaks = df_encode[df_encode['max_conservation'] >= 0.8]
            if not conserved_peaks.empty:
                add_heading_styled(doc, "High-Confidence Conserved Peaks (phastCons >= 0.8)", level=2)
                
                cons_table = doc.add_table(rows=len(conserved_peaks.head(10)) + 1, cols=4)
                cons_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_table_borders(cons_table)
                
                headers = ["Experiment Target", "Accession ID", "Distance to TSS", "Max phastCons"]
                hdr_cells = cons_table.rows[0].cells
                for col_idx, text in enumerate(headers):
                    set_cell_background(hdr_cells[col_idx], "1e293b")
                    set_cell_margins(hdr_cells[col_idx], top=100, bottom=100, left=100, right=100)
                    p = hdr_cells[col_idx].paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    r = p.add_run(text)
                    r.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)
                    
                for r_idx, (_, row) in enumerate(conserved_peaks.head(10).iterrows()):
                    cells = cons_table.rows[r_idx + 1].cells
                    vals = [
                        str(row.get('experiment', 'N/A')),
                        str(row.get('file_accession', 'N/A')),
                        f"{int(row.get('distance_to_tss', 0)):+,} bp",
                        f"{row.get('max_conservation', 0.0):.4f}"
                    ]
                    for col_idx, val in enumerate(vals):
                        set_cell_margins(cells[col_idx], top=60, bottom=60, left=100, right=100)
                        p = cells[col_idx].paragraphs[0]
                        p.paragraph_format.space_after = Pt(2)
                        p.add_run(val)
    else:
        doc.add_paragraph("No UCSC base-by-base evolutionary conservation datasets were available for this window profile.")
        
    # --- 5. GTEx BASELINE TISSUE EXPRESSION PROFILES ---
    add_heading_styled(doc, "5. GTEx Baseline Tissue Profiles", level=1)
    
    if df_gtex is not None and not df_gtex.empty:
        doc.add_paragraph(
            f"Baseline mRNA expression (Transcripts Per Million, TPM) derived from the GTEx (Genotype-Tissue Expression) database "
            f"documents the physiological co-expression levels of the target gene and transcription regulators across normal human tissues."
        )
        
        # Aggregate top tissues and format in table
        try:
            df_pivot = df_gtex.pivot(index='tissue', columns='Gene', values='TPM').reset_index()
            df_pivot['tissue'] = df_pivot['tissue'].apply(lambda x: x.replace('_', ' '))
            df_top_gtex = df_pivot.sort_values(by=gene_name, ascending=False).head(10)
            
            gtex_table = doc.add_table(rows=len(df_top_gtex) + 1, cols=len(df_top_gtex.columns))
            gtex_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_table_borders(gtex_table)
            
            # Write Header
            hdr_cells = gtex_table.rows[0].cells
            for col_idx, col_name in enumerate(df_top_gtex.columns):
                set_cell_background(hdr_cells[col_idx], "475569")
                set_cell_margins(hdr_cells[col_idx], top=100, bottom=100, left=100, right=100)
                p = hdr_cells[col_idx].paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run("Tissue Site" if col_name == "tissue" else col_name)
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                
            # Write rows
            for r_idx, (_, row) in enumerate(df_top_gtex.iterrows()):
                cells = gtex_table.rows[r_idx + 1].cells
                for col_idx, col_name in enumerate(df_top_gtex.columns):
                    val = row[col_name]
                    str_val = f"{val:.2f} TPM" if isinstance(val, (int, float)) else str(val)
                    
                    set_cell_margins(cells[col_idx], top=60, bottom=60, left=100, right=100)
                    p = cells[col_idx].paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    p.add_run(str_val)
        except Exception:
            doc.add_paragraph("An error occurred formatting the baseline expression table.")
    else:
        doc.add_paragraph("No GTEx tissue-specific co-expression median metrics were loaded in the active session.")
        
    # --- 6. STRING EPIGENETIC INTERACTOME ---
    add_heading_styled(doc, "6. STRING Epigenetic Interactome", level=1)
    
    if string_res is not None and string_res.get("partners") is not None and not string_res["partners"].empty:
        doc.add_paragraph(
            "The physical and functional coregulation interactome mapping retrieved from the STRING database shows the "
            "macromolecular complexes surrounding our target factors. These confidence metrics evaluate structural and text-mined associations."
        )
        
        df_part = string_res["partners"].copy()
        df_part = df_part.sort_values(by="score", ascending=False).head(10)
        
        headers = ["Interaction Node A", "Interaction Node B", "Combined score", "Experimental evidence", "Database evidence"]
        string_table = doc.add_table(rows=len(df_part) + 1, cols=5)
        string_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_table_borders(string_table)
        
        hdr_cells = string_table.rows[0].cells
        for col_idx, text in enumerate(headers):
            set_cell_background(hdr_cells[col_idx], "1e293b")
            set_cell_margins(hdr_cells[col_idx], top=100, bottom=100, left=100, right=100)
            p = hdr_cells[col_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
        for r_idx, (_, row) in enumerate(df_part.iterrows()):
            cells = string_table.rows[r_idx + 1].cells
            vals = [
                str(row.get('preferredName_A', 'N/A')),
                str(row.get('preferredName_B', 'N/A')),
                f"{row.get('score', 0.0):.4f}",
                f"{row.get('escore', 0.0):.4f}",
                f"{row.get('dscore', 0.0):.4f}"
            ]
            for col_idx, val in enumerate(vals):
                set_cell_margins(cells[col_idx], top=60, bottom=60, left=100, right=100)
                p = cells[col_idx].paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                p.add_run(val)
    else:
        doc.add_paragraph("No STRING interactome network partner tables were fetched in this analysis session.")
        
    # --- 7. DESIGNED PCR/ChIP PRIMERS ---
    add_heading_styled(doc, "7. Designed PCR / ChIP Primers", level=1)
    
    if primer_results:
        doc.add_paragraph(
            "Below are the optimal primer pairs designed to target overlapping binding peaks or specific promoter regions. "
            "Primer safety evaluations check secondary structures (hairpins, self-dimers, heterodimers) using thermodynamic simulations."
        )
        
        headers = ["Target Label", "Primer Sequence (5' to 3')", "Binding Loc", "Tm (°C)", "GC (%)", "Product Size", "Safety Badge"]
        primer_table = doc.add_table(rows=len(primer_results) * 2 + 1, cols=7)
        primer_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_table_borders(primer_table)
        
        hdr_cells = primer_table.rows[0].cells
        for col_idx, text in enumerate(headers):
            set_cell_background(hdr_cells[col_idx], "475569")
            set_cell_margins(hdr_cells[col_idx], top=100, bottom=100, left=100, right=100)
            p = hdr_cells[col_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
        import tf_explorer.primer_design as pd_tool
        
        r_idx = 1
        for name, res in primer_results:
            if res.get('PRIMER_PAIR_NUM_RETURNED', 0) > 0:
                p = 0
                fp = res[f'PRIMER_LEFT_{p}_SEQUENCE']
                rp = res[f'PRIMER_RIGHT_{p}_SEQUENCE']
                tm_f = res[f'PRIMER_LEFT_{p}_TM']
                tm_r = res[f'PRIMER_RIGHT_{p}_TM']
                prod = res[f'PRIMER_PAIR_{p}_PRODUCT_SIZE']
                
                # Bindings
                fp_idx = res[f'PRIMER_LEFT_{p}'][0]
                rp_idx = res[f'PRIMER_RIGHT_{p}'][0]
                
                # relative location to TSS
                fp_loc = fp_idx - promoter_up
                rp_loc = rp_idx - promoter_up
                
                # Calculate GC content
                def gc_content(seq):
                    return (seq.count('G') + seq.count('C')) / len(seq) * 100.0
                
                gc_f = gc_content(fp)
                gc_r = gc_content(rp)
                
                # Thermodynamic safety
                safety_res = pd_tool.check_primer_safety(fp, rp)
                safety_status = safety_res.get("safety_status", "🟢 Safe")
                
                # Write Forward Row
                cells_f = primer_table.rows[r_idx].cells
                set_cell_background(cells_f[0], "f8fafc")
                vals_f = [
                    f"{name} (FP)",
                    fp,
                    f"{fp_loc:+,} bp",
                    f"{tm_f:.1f}°C",
                    f"{gc_f:.1f}%",
                    f"{prod} bp",
                    safety_status
                ]
                for col_idx, val in enumerate(vals_f):
                    set_cell_margins(cells_f[col_idx], top=60, bottom=60, left=100, right=100)
                    p_cell = cells_f[col_idx].paragraphs[0]
                    p_cell.paragraph_format.space_after = Pt(2)
                    p_cell.add_run(val)
                    
                # Write Reverse Row
                cells_r = primer_table.rows[r_idx + 1].cells
                set_cell_background(cells_r[0], "f8fafc")
                vals_r = [
                    f"{name} (RP)",
                    rp,
                    f"{rp_loc:+,} bp",
                    f"{tm_r:.1f}°C",
                    f"{gc_r:.1f}%",
                    f"{prod} bp",
                    safety_status
                ]
                for col_idx, val in enumerate(vals_r):
                    set_cell_margins(cells_r[col_idx], top=60, bottom=60, left=100, right=100)
                    p_cell = cells_r[col_idx].paragraphs[0]
                    p_cell.paragraph_format.space_after = Pt(2)
                    p_cell.add_run(val)
                    
                r_idx += 2
    else:
        doc.add_paragraph("No specific ChIP or PCR cloning amplicons were designed in this active research session.")
        
    # --- WRITE DOCUMENT ---
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    
    if output_path:
        with open(output_path, "wb") as f:
            f.write(stream.getbuffer())
            
    return stream
